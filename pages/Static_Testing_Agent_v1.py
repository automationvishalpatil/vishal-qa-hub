import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re
import json
from typing import List, Dict, Any, Tuple

# --- Configuration for Static Checks ---
# Keeping the user's additions to MANDATORY_SECTIONS
MANDATORY_SECTIONS = [
    "Introduction",
    "Scope of Work",
    "Assumptions",
    "Acceptance Criteria",
    "Out of Scope",
    "Header/Footer Validation",
    "Mandatory/Optional field"
]

# JSON schema for the LLM's structured output (Used in the agent logic)
GEMINI_ANALYSIS_SCHEMA = {
    "type": "ARRAY",
    "description": "A list of identified defects in the document.",
    "items": {
        "type": "OBJECT",
        "properties": {
            "Defect Type": {"type": "STRING", "description": "e.g., 'Clarity', 'Tone', 'Grammar', 'Consistency'."},
            "Description": {"type": "STRING", "description": "Specific finding and recommendation for correction."},
            "Location": {"type": "STRING", "description": "e.g., 'Section 2.1', 'Paragraph 5', 'Overall Tone'."},
            "Severity": {"type": "STRING", "description": "One of: 'High', 'Medium', 'Low'."}
        },
        "required": ["Defect Type", "Description", "Location", "Severity"]
    }
}

# --- Core Validation Logic ---

def rule_based_analysis(document: Document) -> List[Dict[str, str]]:
    """Performs strict, rule-based structural checks. Returns a list of FAILURES."""
    defects = []
    
    full_text = "\n".join([p.text for p in document.paragraphs])
    
    # 1. Check for Mandatory Sections (Structural Check)
    for section in MANDATORY_SECTIONS:
        # Check if the section name appears as a main heading (case-insensitive)
        if section.upper() not in full_text.upper():
            # Note: The defect type 'Structure/Completeness' is used in analyze_document for consolidation
            defects.append({
                'Defect Type': 'Structure/Completeness',
                'Description': f"Missing mandatory section header: '{section}'. This is a severe structural defect.",
                'Location': 'Document Structure',
                'Severity': 'High'
            })
            
    # 2. Check for Placeholder Keywords (TBD/WIP, etc.)
    PLACEHOLDER_KEYWORDS = [r'\bTBD\b', r'\bTBC\b', r'\bWIP\b', r'\bTODO\b', r'PLACEHOLDER']
    for i, paragraph in enumerate(document.paragraphs):
        for keyword_regex in PLACEHOLDER_KEYWORDS:
            matches = re.findall(keyword_regex, paragraph.text, re.IGNORECASE)
            if matches:
                # Note: The defect type 'Content Placeholder' is used in analyze_document for consolidation
                defects.append({
                    'Defect Type': 'Content Placeholder',
                    'Description': f"Found hardcoded placeholder '{matches[0]}'. Must be resolved before submission.",
                    'Location': f"Paragraph {i + 1}",
                    'Severity': 'Medium'
                })
    
    return defects

def llm_based_analysis(doc_content: str, analysis_goal: str) -> List[Dict[str, str]]:
    """Performs qualitative analysis using the Gemini Agent (simulated)."""
    
    # Placeholder for actual LLM call:
    
    goal_instruction = ""
    if analysis_goal:
        # Pass the goal directly into the system instruction
        goal_instruction = f"Additionally, your primary goal is to focus specifically on this requirement: **{analysis_goal}**."

    # NOTE: The system instruction must be passed to the LLM (if this were a real API call)
    system_instruction = (
        "You are a Senior Technical Writer and Quality Assurance Agent specializing in functional design "
        "documents (FDDs) for telecommunication projects. Your task is to perform a qualitative "
        "static analysis of the document provided. Check for the following: "
        "1. **Clarity and Conciseness:** Is the technical language clear? Are concepts explained well? "
        "2. **Tone:** Is the tone professional, objective, and appropriate for a technical audience? "
        "3. **Consistency:** Are terms, abbreviations, and numerical formats used consistently? "
        "4. **Completeness/Logic:** Are any sections suspiciously brief? Does the logic flow?"
        f"{goal_instruction} Identify and list defects strictly in the required JSON format."
    )
    
    user_query = f"Analyze the following functional document text for qualitative defects and compliance with FDD standards:\n\n---\n{doc_content}\n---"

    # --- SIMULATED RESPONSE (This section is critical for mocking the focused analysis) ---
    
    # Base mock response
    mock_llm_response_json = """
    [
        {
            "Defect Type": "Clarity",
            "Description": "The description of the 'Network_Owner' field in the Scope section is ambiguous. Clarify if it refers to the physical or logical owner.",
            "Location": "Scope of Work, Paragraph 3",
            "Severity": "High"
        },
        {
            "Defect Type": "Tone",
            "Description": "The Conclusion section uses overly casual language ('We think this is fine'). Replace with professional equivalents.",
            "Location": "Conclusion",
            "Severity": "Low"
        },
        {
            "Defect Type": "Consistency",
            "Description": "The acronym 'FDD' is used without definition in the Introduction section.",
            "Location": "Introduction",
            "Severity": "Medium"
        }
    ]
    """
    
    # If a goal is provided, we simulate the LLM finding a related defect
    mock_list = json.loads(mock_llm_response_json)
    if analysis_goal:
        goal_defect = {
            "Defect Type": "Custom Goal Check",
            "Description": f"The Agent focused on the custom goal: '{analysis_goal}'. Found that specific data types (e.g., date formats) are not compliant with internal standards.",
            "Location": "Data Fields Appendix",
            "Severity": "High"
        }
        mock_list.append(goal_defect)

    mock_llm_response_json = json.dumps(mock_list)
    # --- END SIMULATED RESPONSE ---
    
    try:
        llm_defects = json.loads(mock_llm_response_json)
        if not isinstance(llm_defects, list):
            raise ValueError("LLM response did not return a list structure.")
        return llm_defects
    except Exception as e:
        # Return a critical defect object if the LLM call fails
        return [{
            'Defect Type': 'LLM Error (Simulated)',
            'Description': f"Failed to get qualitative analysis: {e}. (Check API key and SDK installation.)",
            'Location': 'LLM Agent',
            'Severity': 'Critical'
        }]

# --- Document Handling and Main Function (Restored to Summary Generation) ---

def analyze_document(doc_file, analysis_goal: str = "") -> Tuple[List[Dict[str, str]], List[Dict[str, str]]]:
    """
    Runs all analyses and formats results into a summary report and detailed defect list.
    Returns: (summary_report_rows, all_defects_found)
    """
    
    # 1. Load the document
    try:
        document = Document(BytesIO(doc_file.read()))
    except Exception as e:
        # Critical failure row for the summary report
        critical_fail_row = {
            'File Name': 'N/A',
            'Pattern': 'File Loading',
            'Structure Element': 'File Loading',
            'Test Result': 'Critical Fail',
            'Remarks': f"Could not process document. Ensure it is a valid .docx file. Error: {e}"
        }
        return [critical_fail_row], []

    full_text = "\n".join([p.text for p in document.paragraphs])
    
    # 2. Run rule-based checks & LLM checks
    rule_defects = rule_based_analysis(document)
    llm_defects = []
    
    if len(full_text) > 50: 
        llm_defects = llm_based_analysis(full_text, analysis_goal)
    else:
        # If content is too short, flag it as a content length issue (which should be handled in the summary)
        llm_defects.append({
            'Defect Type': 'Content Length',
            'Description': 'Document content is too short (< 50 characters) for meaningful analysis.',
            'Location': 'N/A',
            'Severity': 'Medium'
        })


    # 3. Consolidate results into the Summary Report format

    report_rows = []
    
    # --- Check 1: Mandatory Sections (Rule-Based) ---
    missing_sections = [d for d in rule_defects if d['Defect Type'] == 'Structure/Completeness']
    if missing_sections:
        # Extract the missing section names for a concise remark
        remarks = "Missing sections: " + ", ".join([
            d['Description'].split("header: '")[1].split("'. This is")[0] 
            for d in missing_sections
        ])
        report_rows.append({
            'Pattern': 'Document Structure',
            'Structure Element': 'Mandatory Sections',
            'Test Result': 'Fail',
            'Remarks': remarks,
        })
    else:
        report_rows.append({
            'Pattern': 'Document Structure',
            'Structure Element': 'Mandatory Sections',
            'Test Result': 'Pass',
            'Remarks': 'All mandatory sections are present.',
        })

    # --- Check 2: Placeholder Keywords (Rule-Based) ---
    placeholder_defects = [d for d in rule_defects if d['Defect Type'] == 'Content Placeholder']
    if placeholder_defects:
        remarks = f"Found {len(placeholder_defects)} hardcoded placeholders (TBD, WIP, etc.). Fix before finalizing."
        report_rows.append({
            'Pattern': 'Content Quality',
            'Structure Element': 'Placeholder Keywords',
            'Test Result': 'Fail',
            'Remarks': remarks,
        })
    else:
        report_rows.append({
            'Pattern': 'Content Quality',
            'Structure Element': 'Placeholder Keywords',
            'Test Result': 'Pass',
            'Remarks': 'No hardcoded placeholders found.',
        })
        
    # --- Check 3: LLM-based Qualitative Checks (Clarity, Tone, Consistency, Completeness) ---
    # Group all non-goal/non-error LLM findings
    qualitative_defects = [d for d in llm_defects if d['Defect Type'] not in ['Custom Goal Check', 'LLM Error (Simulated)', 'Content Length']]
    
    if any(d['Severity'] in ['High', 'Critical'] for d in qualitative_defects):
        remarks = f"Found {len(qualitative_defects)} qualitative defects. High severity issues detected in Clarity, Tone, or Consistency."
        report_rows.append({
            'Pattern': 'Qualitative Analysis',
            'Structure Element': 'Clarity & Consistency',
            'Test Result': 'Fail',
            'Remarks': remarks,
        })
    elif any(d['Severity'] == 'Medium' for d in qualitative_defects):
         remarks = f"Found {len(qualitative_defects)} qualitative defects. Review Medium severity issues found by the Agent."
         report_rows.append({
            'Pattern': 'Qualitative Analysis',
            'Structure Element': 'Clarity & Consistency',
            'Test Result': 'Pass w/ Warnings',
            'Remarks': remarks,
        })
    elif not any(d['Severity'] in ['Medium', 'High', 'Critical'] for d in qualitative_defects) and not [d for d in llm_defects if d['Defect Type'] == 'Content Length']:
        report_rows.append({
            'Pattern': 'Qualitative Analysis',
            'Structure Element': 'Clarity & Consistency',
            'Test Result': 'Pass',
            'Remarks': 'Qualitative content (clarity, tone, consistency) appears high quality.',
        })

    # --- Check 4: Custom Goal Check (LLM-Based) ---
    custom_goal_defects = [d for d in llm_defects if d['Defect Type'] == 'Custom Goal Check']
    if analysis_goal:
        # Use a short description of the goal for the Structure Element column
        goal_element_name = f'Custom Goal: {analysis_goal[:40]}...' if len(analysis_goal) > 40 else f'Custom Goal: {analysis_goal}'
        
        if custom_goal_defects:
            remarks = custom_goal_defects[0]['Description'] # Use Agent's detailed description
            report_rows.append({
                'Pattern': 'Custom Compliance',
                'Structure Element': goal_element_name,
                'Test Result': 'Fail',
                'Remarks': remarks,
            })
        else:
            report_rows.append({
                'Pattern': 'Custom Compliance',
                'Structure Element': goal_element_name,
                'Test Result': 'Pass',
                'Remarks': f'Document meets the custom analysis goal: "{analysis_goal}".',
            })
            
    # --- Check 5: LLM Error Status ---
    llm_errors = [d for d in llm_defects if d['Defect Type'] == 'LLM Error (Simulated)' or d['Defect Type'] == 'Content Length']
    if llm_errors:
        remarks = llm_errors[0]['Description']
        report_rows.append({
            'Pattern': 'Tool Status',
            'Structure Element': 'LLM Agent',
            'Test Result': 'Critical Fail',
            'Remarks': remarks,
        })
        
    # Return consolidated report rows and all found defects for the detailed table
    all_defects = rule_defects + [d for d in llm_defects if d['Defect Type'] != 'LLM Error (Simulated)' and d['Defect Type'] != 'Content Length']
    return report_rows, all_defects

# --- Streamlit UI (Updated to handle dual return) ---

def main():
    st.set_page_config(
        page_title="Document Static Tester (Agent Enhanced)", 
        layout="wide", 
        initial_sidebar_state="expanded"
    )

    st.title("📄 Functional Document Static Tester (Agent Enhanced)")
    st.markdown("This tool generates a **Summary Validation Report** by combining **Rule-Based Checks** with a **Gemini Agent** (simulated) for qualitative analysis.")

    # File Uploader
    uploaded_file = st.file_uploader(
        "Upload Word Document (.docx)", 
        type=["docx"], 
        accept_multiple_files=False
    )
    
    # New Goal Input
    analysis_goal = st.text_area(
        "💡 Custom Analysis Goal / Focus (Optional)",
        placeholder="E.g., Ensure all feature descriptions are suitable for a non-technical audience or Verify compliance with IEEE standard 1059.",
        height=100
    )
    
    # Submit Button
    if uploaded_file is not None:
        if st.button("Submit Document for Validation", type="primary"):
            with st.spinner("Analyzing document for static and qualitative defects using the Agent..."):
                
                # Get the filename for the report
                file_name_display = uploaded_file.name
                
                # Call the analysis function to get both the summary report and the detailed defects
                report_summary, detail_defects = analyze_document(uploaded_file, analysis_goal)
                
                # --- 1. Generate Summary Validation Report (User's requested format) ---
                
                # Add 'File Name' to all rows and define final columns order
                for row in report_summary:
                    row['File Name'] = file_name_display
                
                summary_cols = ['File Name', 'Pattern', 'Structure Element', 'Test Result', 'Remarks']
                df_summary = pd.DataFrame(report_summary)[summary_cols]
                
                st.subheader("✅ Summary Validation Report")
                
                # Define function to color the 'Test Result' column
                def color_result(val):
                    color_map = {
                        'Pass': 'background-color: #d4edda; color: #155724; font-weight: bold;',
                        'Pass w/ Warnings': 'background-color: #fff3cd; color: #856404; font-weight: bold;',
                        'Fail': 'background-color: #f8d7da; color: #721c24; font-weight: bold;',
                        'Critical Fail': 'background-color: #f5c6cb; color: #721c24; font-weight: bold;',
                    }
                    return color_map.get(val, '')

                # Display the summary report
                st.dataframe(
                    df_summary.style.applymap(color_result, subset=['Test Result']), 
                    use_container_width=True, 
                    hide_index=True
                )
                
                # --- 2. Generate Detailed Defect List (Original format) ---
                
                st.subheader("🐞 Detailed Defect Findings (Agent Output)")
                
                # Convert the detailed defects into a pandas DataFrame
                df_details = pd.DataFrame(detail_defects)
                
                # Sort the defects by severity for the detailed view
                severity_order = {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1, 'Success': 0}
                if 'Severity' in df_details.columns:
                    df_details['Sort_Severity'] = df_details['Severity'].map(severity_order)
                    df_details = df_details.sort_values(by='Sort_Severity', ascending=False).drop(columns=['Sort_Severity'])
                
                # Display the results table
                if not df_details.empty:
                    st.dataframe(df_details, use_container_width=True, hide_index=True)
                else:
                    st.info("The Agent found no detailed defects of High, Medium, or Low severity.")

                # Final Summary
                fail_count = len(df_summary[df_summary['Test Result'].isin(['Fail', 'Critical Fail'])])
                
                if fail_count > 0:
                    st.error(f"⚠️ **{fail_count}** critical compliance checks failed. Review the Summary Validation Report for immediate attention items.")
                elif 'Pass w/ Warnings' in df_summary['Test Result'].unique():
                    st.warning("Document passed all critical checks but the Agent issued warnings for qualitative defects.")
                else:
                    st.balloons()
                    st.success("Document achieved a clean 'Pass' status on all static and qualitative checks!")

    st.sidebar.header("Check Configuration")
    st.sidebar.markdown("**Mandatory Sections (Rule-Based):**")
    st.sidebar.code("\n".join(MANDATORY_SECTIONS))
    st.sidebar.markdown("""
    ---
    ### 🧠 LLM Agent Checks
    The Agent performs qualitative checks for:
    - Clarity & Conciseness
    - Professional Tone
    - Consistency
    - Logical Completeness
    """)
    if analysis_goal:
        st.sidebar.markdown(f"**Current Goal Focus:**\n_{analysis_goal}_")


if __name__ == "__main__":
    main()