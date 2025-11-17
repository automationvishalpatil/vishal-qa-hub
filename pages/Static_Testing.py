import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re

# --- Configuration for Static Checks ---
MANDATORY_SECTIONS = [
    "Introduction",
    "Scope of Work",
    "Assumptions",
    "Acceptance Criteria",
    "Out of Scope"
]

PLACEHOLDER_KEYWORDS = [
    r'\bTBD\b', # To Be Determined
    r'\bTBC\b', # To Be Confirmed
    r'\bWIP\b', # Work in Progress
    r'\bTODO\b',
    r'PLACEHOLDER',
    r'INSERT TEXT HERE'
]

# --- Core Validation Logic ---

def analyze_document(doc_file):
    """
    Reads the Word document content and performs static analysis checks.

    Args:
        doc_file (UploadedFile): The file object from st.file_uploader.

    Returns:
        list: A list of dictionaries, where each dictionary is a defect.
    """
    defects = []
    
    # Load the document from the uploaded file's bytes
    try:
        document = Document(BytesIO(doc_file.read()))
    except Exception as e:
        defects.append({
            'Defect Type': 'File Error',
            'Description': f"Could not process document. Ensure it is a valid .docx file. Error: {e}",
            'Location': 'N/A',
            'Severity': 'Critical'
        })
        return defects

    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)
    
    # 1. Check for Mandatory Sections
    doc_content_string = "\n".join(full_text)
    
    for section in MANDATORY_SECTIONS:
        # Check if the section name appears as a main heading
        if section.upper() not in doc_content_string.upper():
            defects.append({
                'Defect Type': 'Structure/Completeness',
                'Description': f"Missing mandatory section header: '{section}'",
                'Location': 'Document Structure',
                'Severity': 'High'
            })

    # 2. Check for Placeholder Keywords
    for i, paragraph in enumerate(document.paragraphs):
        for keyword_regex in PLACEHOLDER_KEYWORDS:
            # Use regex to find case-insensitive matches
            matches = re.findall(keyword_regex, paragraph.text, re.IGNORECASE)
            
            if matches:
                defects.append({
                    'Defect Type': 'Content Placeholder',
                    'Description': f"Placeholder found: '{matches[0]}'. Text requires update.",
                    'Location': f"Paragraph {i + 1}",
                    'Severity': 'Medium'
                })

    # 3. Check for Readability (e.g., excessively long sentences or paragraphs)
    # A paragraph is too long if it exceeds 150 words (an arbitrary rule for functional docs)
    MAX_WORDS = 150
    for i, paragraph in enumerate(document.paragraphs):
        word_count = len(paragraph.text.split())
        if word_count > MAX_WORDS:
            defects.append({
                'Defect Type': 'Readability',
                'Description': f"Paragraph contains {word_count} words (limit is {MAX_WORDS}). Consider splitting.",
                'Location': f"Paragraph {i + 1}",
                'Severity': 'Low'
            })

    # If no defects found
    if not defects:
         defects.append({
            'Defect Type': 'Quality Status',
            'Description': 'No static defects found based on current checks!',
            'Location': 'N/A',
            'Severity': 'Success'
        })

    return defects

# --- Streamlit UI ---

def main():
    st.set_page_config(
        page_title="Document Static Testing", 
        layout="wide", 
        initial_sidebar_state="expanded"
    )

    st.title("📄 Functional Document Static Testing")
    st.markdown("Upload your functional design document (.docx) to check for structural and content defects.")

    # File Uploader
    uploaded_file = st.file_uploader(
        "Upload Word Document (.docx)", 
        type=["docx"], 
        accept_multiple_files=False
    )
    
    # Submit Button
    if uploaded_file is not None:
        if st.button("Submit Document for Validation", type="primary"):
            with st.spinner("Analyzing document for static defects..."):
                defect_list = analyze_document(uploaded_file)
                
                # Convert the results into a pandas DataFrame for structured display
                df_defects = pd.DataFrame(defect_list)
                
                # Sort the defects by severity
                severity_order = {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1, 'Success': 0}
                df_defects['Sort_Severity'] = df_defects['Severity'].map(severity_order)
                df_defects = df_defects.sort_values(by='Sort_Severity', ascending=False).drop(columns=['Sort_Severity'])
                
                st.subheader("📊 Validation Results")
                
                # Display the results table
                st.dataframe(df_defects, use_container_width=True, hide_index=True)
                
                # Summary
                total_defects = len(df_defects[df_defects['Severity'] != 'Success'])
                if total_defects == 0:
                    st.balloons()
                    st.success("Document looks clean! No static defects detected.")
                else:
                    st.warning(f"Found **{total_defects}** defects that require attention.")

    st.sidebar.header("Check Configuration")
    st.sidebar.markdown("**Mandatory Sections:**")
    st.sidebar.code("\n".join(MANDATORY_SECTIONS))
    st.sidebar.markdown("**Placeholder Keywords Checked:**")
    st.sidebar.code("\n".join([r.strip(r'\b') for r in PLACEHOLDER_KEYWORDS]))


if __name__ == "__main__":
    main()