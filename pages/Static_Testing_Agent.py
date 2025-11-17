import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re
import json
from typing import List, Dict, Any

# --- Configuration for Static Checks ---
MANDATORY_SECTIONS = [
    "Introduction",
    "Scope of Work",
    "Assumptions",
    "Acceptance Criteria",
    "Out of Scope"
]

# JSON schema for the LLM's structured output
GEMINI_ANALYSIS_SCHEMA = {
    "type": "ARRAY",
    "description": "A list of identified defects in the document.",
    "items": {
        "type": "OBJECT",
        "properties": {
            "Defect Type": {"type": "STRING", "description": "e.g., 'Clarity', 'Tone', 'Grammar', 'Consistency'."},
            "Description": {"type": "STRING", "description": "Specific finding and recommendation for correction."},
            "Location": {"type": "STRING", "description": "e.g., 'Section 2.1', 'Paragraph 5', 'Overall Tone'."},
            "Severity": {"type": "STRING", "description": "One of: 'High', 'Medium', 'Low'."}
        },
        "required": ["Defect Type", "Description", "Location", "Severity"]
    }
}

# --- Core Validation Logic ---

def rule_based_analysis(document: Document) -> List[Dict[str, str]]:
    """Performs strict, rule-based structural checks."""
    defects = []
    
    full_text = "\n".join([p.text for p in document.paragraphs])
    
    # 1. Check for Mandatory Sections (Structural Check)
    for section in MANDATORY_SECTIONS:
        # Check if the section name appears as a main heading (case-insensitive)
        if section.upper() not in full_text.upper():
            defects.append({
                'Defect Type': 'Structure/Completeness',
                'Description': f"Missing mandatory section header: '{section}'. This is a severe structural defect.",
                'Location': 'Document Structure',
                'Severity': 'High'
            })
            
    # 2. Check for Placeholder Keywords (TBD/WIP, etc.)
    # The LLM can handle qualitative placeholders, but this ensures no literal placeholders remain.
    PLACEHOLDER_KEYWORDS = [r'\bTBD\b', r'\bTBC\b', r'\bWIP\b', r'\bTODO\b', r'PLACEHOLDER']
    for i, paragraph in enumerate(document.paragraphs):
        for keyword_regex in PLACEHOLDER_KEYWORDS:
            matches = re.findall(keyword_regex, paragraph.text, re.IGNORECASE)
            if matches:
                defects.append({
                    'Defect Type': 'Content Placeholder',
                    'Description': f"Found hardcoded placeholder '{matches[0]}'. Must be resolved before submission.",
                    'Location': f"Paragraph {i + 1}",
                    'Severity': 'Medium'
                })
    
    return defects

def llm_based_analysis(doc_content: str) -> List[Dict[str, str]]:
    """
    Performs qualitative analysis using the Gemini Agent.
    
    NOTE: To make this function operational, you must install the Google GenAI SDK 
    ('pip install google-genai') and provide a valid API key (e.g., via st.secrets).
    """
    
    # --- LLM API INTEGRATION ---
    
    # Placeholder/Instructions for actual LLM call:
    # 1. Initialize the client: client = GeminiClient(api_key=st.secrets["GEMINI_API_KEY"])
    # 2. Define the system instruction and query.
    # 3. Call the API using the defined JSON schema.
    
    system_instruction = (
        "You are a Senior Technical Writer and Quality Assurance Agent specializing in functional design "
        "documents (FDDs) for telecommunication projects. Your task is to perform a qualitative "
        "static analysis of the document provided. Check for the following: "
        "1. **Clarity and Conciseness:** Is the technical language clear? Are concepts explained well? "
        "2. **Tone:** Is the tone professional, objective, and appropriate for a technical audience? "
        "3. **Consistency:** Are terms, abbreviations, and numerical formats used consistently? "
        "4. **Completeness/Logic:** Are any sections suspiciously brief? Does the logic flow?"
        "Identify and list defects strictly in the required JSON format."
    )
    
    user_query = f"Analyze the following functional document text for qualitative defects and compliance with FDD standards:\n\n---\n{doc_content}\n---"

    # --- SIMULATED RESPONSE (Replace with actual API call) ---
    
    # In a real implementation, you would call the model here.
    # Example (MOCKED JSON RESPONSE):
    
    mock_llm_response_json = """
    [
        {
            "Defect Type": "Clarity",
            "Description": "The description of the 'Network_Owner' field in the Scope section is ambiguous. Clarify if it refers to the physical or logical owner.",
            "Location": "Scope of Work, Paragraph 3",
            "Severity": "High"
        },
        {
            "Defect Type": "Tone",
            "Description": "The Conclusion section uses overly casual language ('We think this is fine'). Replace with professional equivalents.",
            "Location": "Conclusion",
            "Severity": "Low"
        },
        {
            "Defect Type": "Completeness",
            "Description": "The 'Assumptions' section is only one sentence long. A document of this scope requires at least 3-5 explicit, detailed assumptions regarding systems or external dependencies.",
            "Location": "Assumptions Section",
            "Severity": "Medium"
        }
    ]
    """
    
    # --- END SIMULATED RESPONSE ---
    
    try:
        # In a real app, 'response = client.generate_content(...)' would return the JSON string.
        llm_defects = json.loads(mock_llm_response_json)
        # Validate that the LLM returned a list of dictionaries as expected by the schema
        if not isinstance(llm_defects, list):
            raise ValueError("LLM response did not return a list structure.")
        return llm_defects
    except Exception as e:
        return [{
            'Defect Type': 'LLM Error (Simulated)',
            'Description': f"Failed to get qualitative analysis: {e}. (Check API key and SDK installation.)",
            'Location': 'LLM Agent',
            'Severity': 'Critical'
        }]

# --- Document Handling and Main Function ---

def analyze_document(doc_file) -> List[Dict[str, str]]:
    """Master function to run all analyses."""
    defects = []
    
    # Load the document
    try:
        document = Document(BytesIO(doc_file.read()))
    except Exception as e:
        defects.append({
            'Defect Type': 'File Error',
            'Description': f"Could not process document. Ensure it is a valid .docx file. Error: {e}",
            'Location': 'N/A',
            'Severity': 'Critical'
        })
        return defects
    
    # Extract full content for LLM
    full_text = "\n".join([p.text for p in document.paragraphs])

    # 1. Run Rule-Based Checks
    defects.extend(rule_based_analysis(document))
    
    # 2. Run LLM-Based Checks
    # Only run LLM if the text isn't trivially empty
    if len(full_text) > 50: 
        defects.extend(llm_based_analysis(full_text))
    else:
        defects.append({
            'Defect Type': 'Content Error',
            'Description': 'Document content is too short to perform meaningful LLM analysis.',
            'Location': 'N/A',
            'Severity': 'Medium'
        })

    # If all checks pass
    if all(d.get('Severity') != 'High' and d.get('Severity') != 'Critical' for d in defects):
        defects.append({
            'Defect Type': 'Quality Status',
            'Description': 'No critical or high-severity defects found by the agent or rules!',
            'Location': 'N/A',
            'Severity': 'Success'
        })
        # Remove the placeholder success check from the previous version if other defects exist
        defects = [d for d in defects if d.get('Severity') != 'Success']
        if not defects:
             defects.append({'Defect Type': 'Quality Status', 'Description': 'Perfect! No defects found.', 'Location': 'N/A', 'Severity': 'Success'})


    return defects

# --- Streamlit UI ---

def main():
    st.set_page_config(
        page_title="Document Static Testing Agent", 
        layout="wide", 
        initial_sidebar_state="expanded"
    )

    st.title("📄 Functional Document Static Testing Agent")
    st.markdown("This tool combines **Rule-Based Checks** with a **Gemini Agent** (simulated) for qualitative analysis.")

    # File Uploader
    uploaded_file = st.file_uploader(
        "Upload Word Document (.docx)", 
        type=["docx"], 
        accept_multiple_files=False
    )
    
    # Submit Button
    if uploaded_file is not None:
        if st.button("Submit Document for Validation", type="primary"):
            with st.spinner("Analyzing document for static and qualitative defects using the Agent..."):
                defect_list = analyze_document(uploaded_file)
                
                # Convert the results into a pandas DataFrame for structured display
                df_defects = pd.DataFrame(defect_list)
                
                # Sort the defects by severity
                severity_order = {'Critical': 4, 'High': 3, 'Medium': 2, 'Low': 1, 'Success': 0}
                df_defects['Sort_Severity'] = df_defects['Severity'].map(severity_order)
                df_defects = df_defects.sort_values(by='Sort_Severity', ascending=False).drop(columns=['Sort_Severity'])
                
                st.subheader("📊 Validation Results")
                
                # Display the results table
                st.dataframe(df_defects, use_container_width=True, hide_index=True)
                
                # Summary
                total_defects = len(df_defects[df_defects['Severity'] != 'Success'])
                if total_defects == 0:
                    st.balloons()
                    st.success("Document looks clean! No static defects detected.")
                else:
                    st.warning(f"Found **{total_defects}** defects that require attention. (Includes qualitative feedback from agent.)")

    st.sidebar.header("Check Configuration")
    st.sidebar.markdown("**Mandatory Sections (Rule-Based):**")
    st.sidebar.code("\n".join(MANDATORY_SECTIONS))
    st.sidebar.markdown("""
    ---
    ### 🧠 LLM Agent Checks
    The Agent performs qualitative checks for:
    - Clarity & Conciseness
    - Professional Tone
    - Consistency
    - Logical Completeness
    Goal: Generate a summary table to validate the structure of Aged Debt Report Instructions: Extract and organise validation results in a tabular format, grouped by Feed Name or Report Name. For each feed, validate the following structure elements: •	Presence and correctness of Header & Footer rows •	Field types (do not infer types from similar fields in other feeds) •	Mandatory/Optional status of each field •	Format compliance •	Field Description The response table should include the following columns: •	File Name Format •	Structure Element •	Test Result •	Remarks Expectations: The Remarks column must clearly describe any issues found, including specific details of the identified problem, documentation reference with Headings & page number along with list of affected Field Names. Test Result column should have only 3 values - 'Pass', 'Fail', 'Potential defect' Output Format: Markdown table or structured tabular format.
    """)


if __name__ == "__main__":
    main()